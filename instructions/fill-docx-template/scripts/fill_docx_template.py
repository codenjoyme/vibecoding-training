#!/usr/bin/env python3
"""
fill_docx_template.py — Fill a DOCX template with values from a .properties file.

The original template is never modified: the script copies it to --output, then
opens the copy and replaces every <key> placeholder with the matching value.

Usage:
    python fill_docx_template.py --template Template.docx \
                                  --properties data.properties \
                                  --output Filled.docx
"""

import argparse
import shutil
import sys
from pathlib import Path

from docx import Document


def load_properties(path: str) -> dict:
    """Return {key: value} from a standard key=value properties file."""
    props = {}
    with open(path, encoding="utf-8") as fh:
        for lineno, raw in enumerate(fh, 1):
            line = raw.strip()
            if not line or line.startswith("#"):
                continue
            if "=" not in line:
                print(
                    f"Warning: line {lineno} skipped (no '=' found): {line!r}",
                    file=sys.stderr,
                )
                continue
            key, _, value = line.partition("=")
            props[key.strip()] = value.strip()
    return props


def _replace_in_paragraph(para, props: dict) -> None:
    """Replace <key> placeholders in a paragraph's runs.

    Two-phase strategy:
    1. Fast path — replace within each individual run (preserves per-run formatting).
    2. Slow path — if a placeholder is still present (split across runs), merge all
       runs into the first one and replace there (minor formatting loss for that run).
    """
    # Quick check: does this paragraph contain any placeholder at all?
    full_text = "".join(run.text for run in para.runs)
    if not any(f"<{k}>" in full_text for k in props):
        return

    # Phase 1: in-place per-run replacement
    for run in para.runs:
        for key, value in props.items():
            placeholder = f"<{key}>"
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)

    # Check whether any placeholder survived (split runs case)
    remaining = "".join(run.text for run in para.runs)
    if not any(f"<{k}>" in remaining for k in props):
        return

    # Phase 2: merge all runs into the first, then replace
    new_text = remaining
    for key, value in props.items():
        new_text = new_text.replace(f"<{key}>", value)

    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""


def fill_template(template_path: str, properties_path: str, output_path: str) -> None:
    """Copy template → output, then replace all <key> placeholders."""
    template = Path(template_path)
    output = Path(output_path)

    if not template.exists():
        print(f"Error: template not found: {template}", file=sys.stderr)
        sys.exit(1)

    if not Path(properties_path).exists():
        print(f"Error: properties file not found: {properties_path}", file=sys.stderr)
        sys.exit(1)

    props = load_properties(properties_path)
    if not props:
        print(
            "Warning: properties file is empty — output will be identical to template.",
            file=sys.stderr,
        )

    # Copy template first (original stays untouched)
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, output)

    # Open the copy and replace placeholders
    doc = Document(str(output))

    # Body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, props)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, props)

    doc.save(str(output))
    print(f"Output written to: {output}")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Fill a DOCX template with values from a .properties file."
    )
    parser.add_argument(
        "--template",
        required=True,
        metavar="TEMPLATE.docx",
        help="Path to the source .docx template (never modified).",
    )
    parser.add_argument(
        "--properties",
        required=True,
        metavar="data.properties",
        help="Path to the .properties file (key=value, one per line).",
    )
    parser.add_argument(
        "--output",
        required=True,
        metavar="OUTPUT.docx",
        help="Destination path for the filled .docx file.",
    )
    args = parser.parse_args()
    fill_template(args.template, args.properties, args.output)


if __name__ == "__main__":
    main()
